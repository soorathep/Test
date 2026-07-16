#!/usr/bin/env python3
"""
Turn returned member forms into website data.

The loop this is built for
--------------------------
1. Send everyone `member_form.docx`.
2. They send back the filled form plus one photo.
3. Drop every file you get into `inbox/` — forms and photos together, any names.
4. Run this. It writes `_data/people.yml` and square photos in `img/people/`.
5. Read the YAML, fix anything odd, commit.

    pip install python-docx pyyaml pillow
    python tools/make_people.py

It never overwrites `_data/people.yml` without showing you the diff first, and it
never invents a value: a field left blank in the form is left out of the YAML.

Photos: phone photos arrive at 3000 px in every aspect ratio. Each is cropped to
a square biased toward the upper third (where faces sit) and saved at 600 px.
The site renders them grayscale until hover, which is what makes a set of
mismatched phone snaps still look like one group.
"""
import re, sys, unicodedata, pathlib, difflib
import yaml
from docx import Document
from PIL import Image, ImageOps

ROOT = pathlib.Path(__file__).resolve().parent.parent
INBOX = ROOT / "inbox"
OUT_PHOTOS = ROOT / "img" / "people"
OUT_YAML = ROOT / "_data" / "people.yml"
PHOTO_PX = 600

# form label  ->  yaml key. Matching is fuzzy, so a member who retypes
# "Full Name" or adds a colon still parses.
FIELDS = {
    "full name": "name",
    "short name": "short",
    "role": "role",
    "year joined": "since",
    "research, in one sentence": "topic",
    "previous degree": "prev",
    "orcid or google scholar": "link",
}
HINT = re.compile(r"^(exactly as|optional|phd student /|e\.g\.|max 20 words)", re.I)


def slugify(name):
    s = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode()
    s = re.sub(r"[^a-zA-Z0-9]+", "-", s).strip("-").lower()
    return "-".join(s.split("-")[:2]) or "member"


def clean(v):
    v = " ".join(v.split())
    if not v or HINT.match(v):        # they left the grey hint text in place
        return ""
    return v


def read_form(path):
    doc = Document(path)
    if not doc.tables:
        print(f"  !! {path.name}: no table found — is this the right form?")
        return None
    rec = {}
    for row in doc.tables[0].rows:
        label = " ".join(row.cells[0].text.split("\n")[0].split()).strip(" :*").lower()
        key = FIELDS.get(label)
        if key is None:                       # fuzzy fallback
            m = difflib.get_close_matches(label, FIELDS, n=1, cutoff=0.75)
            key = FIELDS[m[0]] if m else None
        if key is None:
            continue
        val = clean(row.cells[1].text)
        if val:
            rec[key] = val
    if not rec.get("name"):
        print(f"  !! {path.name}: no name filled in — skipped")
        return None
    if rec.get("since"):
        m = re.search(r"(20\d\d)", rec["since"])
        rec["since"] = int(m.group(1)) if m else rec.pop("since")
    return rec


def square(src, dest):
    im = Image.open(src)
    im = ImageOps.exif_transpose(im).convert("RGB")   # phones lie about rotation
    w, h = im.size
    side = min(w, h)
    left = (w - side) // 2
    top = int((h - side) * 0.25)                      # bias up: faces are not centred
    im = im.crop((left, top, left + side, top + side)).resize((PHOTO_PX, PHOTO_PX), Image.LANCZOS)
    im.save(dest, "JPEG", quality=86, optimize=True, progressive=True)


def main():
    if not INBOX.exists():
        INBOX.mkdir()
        print(f"Created {INBOX.relative_to(ROOT)}/ — put the returned forms and photos in it, then run again.")
        return
    forms = sorted(p for p in INBOX.iterdir() if p.suffix.lower() == ".docx" and not p.name.startswith("~$"))
    photos = [p for p in INBOX.iterdir() if p.suffix.lower() in {".jpg", ".jpeg", ".png", ".heic", ".webp"}]
    if not forms:
        print(f"No .docx forms in {INBOX.relative_to(ROOT)}/")
        return
    OUT_PHOTOS.mkdir(parents=True, exist_ok=True)

    people, used = [], set()
    print(f"{len(forms)} forms, {len(photos)} photos\n")
    for f in forms:
        rec = read_form(f)
        if not rec:
            continue
        rec["slug"] = slugify(rec["name"])

        # match a photo: by form filename, then by the person's name
        keys = [f.stem.lower()] + [w.lower() for w in re.split(r"[\s,]+", rec["name"]) if len(w) > 2]
        hit = next((p for p in photos if p not in used
                    and any(k in p.stem.lower() or p.stem.lower() in k for k in keys)), None)
        if hit:
            used.add(hit)
            dest = OUT_PHOTOS / f"{rec['slug']}.jpg"
            try:
                square(hit, dest)
                rec["photo"] = dest.name
                note = f"photo {hit.name}"
            except Exception as e:
                note = f"PHOTO FAILED ({e}) — convert it to JPEG by hand"
        else:
            note = "NO PHOTO MATCHED — rename the photo to include their name"

        order = ["slug", "name", "short", "role", "since", "topic", "prev", "photo", "link"]
        people.append({k: rec[k] for k in order if k in rec})
        print(f"  {rec['name']:<34} {rec.get('role','?'):<18} {note}")

    orphan = [p.name for p in photos if p not in used]
    if orphan:
        print("\n  photos with no form:", ", ".join(orphan))

    body = ("# Current lab members. Generated by tools/make_people.py from the returned\n"
            "# member_form.docx files. Hand-editing is fine — it is a plain list.\n"
            "# Re-running the script rewrites this file.\n\n"
            + yaml.dump(people, allow_unicode=True, sort_keys=False, width=100))

    if OUT_YAML.exists() and OUT_YAML.read_text(encoding="utf-8") != body:
        print(f"\n{OUT_YAML.relative_to(ROOT)} would change:")
        old = OUT_YAML.read_text(encoding="utf-8").splitlines()
        for line in list(difflib.unified_diff(old, body.splitlines(), "current", "new", lineterm=""))[:40]:
            print("   ", line)
        if input("\nWrite it? [y/N] ").strip().lower() != "y":
            print("Nothing written.")
            return
    OUT_YAML.write_text(body, encoding="utf-8")
    print(f"\nWrote {OUT_YAML.relative_to(ROOT)} ({len(people)} people) and {OUT_PHOTOS.relative_to(ROOT)}/")
    print("Read the YAML before committing. Anything blank in a form is simply absent — that is intended.")


if __name__ == "__main__":
    main()
